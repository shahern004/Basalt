"""One-time script: convert MP template placeholders to Jinja2 syntax for docxtpl.

Opens the original MP .docx template, merges XML runs that split placeholders,
converts {curly-brace} and [bracket] placeholders to {{ jinja2 }} tags, and
writes the result to templates/MP.docx.

Usage:
    cd basalt-stack/tools/rmf-generator
    python retag_template.py [--src PATH] [--dst PATH] [--verify]
"""

import argparse
import sys
from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph


# ─── Paths (relative to this script) ─────────────────────────────────────────

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent.parent.parent
DEFAULT_SRC = (
    REPO_ROOT / "rmf-plan-templates" / "MP-Media Protection Plan - TemplateRev5.docx"
)
DEFAULT_DST = SCRIPT_DIR / "templates" / "MP.docx"


# ─── Placeholder → Jinja2 Mappings ───────────────────────────────────────────
# These dicts are the authoritative source of truth for placeholder-to-variable-
# name translation. B3 decision: the LLM schema will return Jinja2 variable names
# directly (e.g. "odp_defined_frequency"), so no runtime mapping layer is needed.

# Simple 1:1 mappings — every occurrence maps to the same Jinja2 variable.
SIMPLE_MAPPINGS = {
    # --- System metadata (curly braces) ---
    "{Organization}": "{{ organization }}",
    "(Organization}": "{{ organization }}",  # typo in "Prepared by:" para
    "{System Name}": "{{ system_name }}",
    "{System Acronym}": "{{ system_acronym }}",
    "{optional: Higher Level Organization}": "{{ higher_level_organization }}",
    "{Logo}": "{{ logo }}",
    "{appropriate entities}": "{{ distribution_entities }}",
    "{appropriate office}": "{{ distribution_office }}",
    # --- ODP brackets (short form) ---
    "[defined frequency]": "{{ odp_defined_frequency }}",
    "[defined events]": "{{ odp_defined_events }}",
    "[defined roles]": "{{ odp_defined_roles }}",
    "[defined official]": "{{ odp_defined_official }}",
    "[defined types of digital and/or non-digital media]": "{{ odp_types_digital_nondigital_media }}",
    "[defined types of media exempted from marking]": "{{ odp_types_media_exempted_marking }}",
    "[defined controlled areas]": "{{ odp_controlled_areas }}",
    "[defined types of system media]": "{{ odp_types_system_media }}",
    "[defined automated mechanisms]": "{{ odp_automated_mechanisms }}",
    "[defined controls]": "{{ odp_controls }}",
    "[defined system media]": "{{ odp_system_media }}",
    "[defined systems/components]": "{{ odp_systems_components }}",
    "[defined circumstances]": "{{ odp_circumstances }}",
    "[defined sanitization techniques and procedures]": "{{ odp_sanitization_techniques }}",
    "[defined system media downgrading process]": "{{ odp_media_downgrading_process }}",
    "[defined system media requiring downgrading]": "{{ odp_media_requiring_downgrading }}",
    "[defined conditions]": "{{ odp_conditions }}",
    "[restrict, prohibit]": "{{ odp_restrict_prohibit }}",
    # --- Nested bracket (entire outer bracket = one ODP) ---
    "[remotely, under [defined conditions]]": "{{ odp_remote_purge_conditions }}",
    # --- Assignment variants (long form → same variable as short form) ---
    # \xa0 = non-breaking space used in the template
    "[Assignment (one or more):\xa0organization-level,"
    "\xa0mission/business process-level,"
    "\xa0system-level]": "{{ odp_policy_level }}",
    "[Assignment: organization-defined types of digital and/or non-digital media]": "{{ odp_types_digital_nondigital_media }}",
    "[Assignment: organization-defined controlled areas]": "{{ odp_controlled_areas }}",
    "[Assignment: organization-defined types of system media]": "{{ odp_types_system_media }}",
}

# Positional mappings — different Jinja2 variable for each occurrence (ordered).
# The Nth occurrence in document order maps to the Nth entry in the list.
POSITIONAL_MAPPINGS = {
    "{Low, Moderate, or High}": [
        "{{ confidentiality }}",  # 1st: Confidentiality line
        "{{ integrity }}",       # 2nd: Integrity line
        "{{ availability }}",    # 3rd: Availability line
    ],
    "{Name, Rank, Organization}": [
        "{{ ao_name }}",         # 1st: Authorizing Official block
        "{{ issm_name }}",       # 2nd: ISSM block
    ],
    "{Title}": [
        "{{ ao_title }}",        # 1st: AO block
        "{{ issm_title }}",      # 2nd: ISSM block
    ],
}


# ─── Core Logic ───────────────────────────────────────────────────────────────

def get_all_paragraphs(doc: DocxDocument) -> Iterator[Paragraph]:
    """Yield all paragraphs in document order: body paragraphs then table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def has_any_placeholder(text: str) -> bool:
    """Return True if text contains any known placeholder pattern."""
    all_keys = (*SIMPLE_MAPPINGS, *POSITIONAL_MAPPINGS)
    return any(key in text for key in all_keys)


def apply_replacements(text: str, occurrence_counter: dict[str, int]) -> str:
    """Apply all placeholder→Jinja2 replacements to text.

    Order matters:
    1. Positional replacements first (one occurrence at a time, left-to-right).
    2. Simple replacements sorted longest-first so nested brackets like
       [remotely, under [defined conditions]] are matched before [defined conditions].
    """
    # 1. Positional replacements
    for key, values in POSITIONAL_MAPPINGS.items():
        while key in text:
            idx = occurrence_counter.get(key, 0)
            if idx >= len(values):
                raise ValueError(
                    f"Unexpected occurrence #{idx + 1} of positional placeholder "
                    f"'{key}' (only {len(values)} mappings defined)"
                )
            replacement = values[idx]
            occurrence_counter[key] = idx + 1
            text = text.replace(key, replacement, 1)

    # 2. Simple replacements (longest key first)
    for key in sorted(SIMPLE_MAPPINGS, key=len, reverse=True):
        if key in text:
            text = text.replace(key, SIMPLE_MAPPINGS[key])

    return text


def process_paragraph(paragraph: Paragraph, occurrence_counter: dict[str, int]) -> bool:
    """Merge split runs and apply Jinja2 replacements to one paragraph.

    Strategy: join all run texts, apply replacements to the joined string,
    then put the result into run[0] and clear subsequent runs. This preserves
    run[0]'s XML formatting (bold, font, size, etc.) for the whole paragraph.
    Empty runs are invisible in Word and don't affect rendering.
    """
    if not paragraph.runs:
        return False

    full_text = "".join(run.text for run in paragraph.runs)

    if not has_any_placeholder(full_text):
        return False

    new_text = apply_replacements(full_text, occurrence_counter)

    # Collapse into first run, clear the rest
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""

    return True


def retag(src_path: Path, dst_path: Path) -> tuple[int, dict[str, int]]:
    """Open template, re-tag all placeholders, save to dst_path."""
    doc = Document(str(src_path))
    occurrence_counter = {}
    modified_count = 0

    for para in get_all_paragraphs(doc):
        if process_paragraph(para, occurrence_counter):
            modified_count += 1

    Path(dst_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_path))

    return modified_count, occurrence_counter


def verify(dst_path: Path) -> set[str]:
    """Print Jinja2 variables detected by docxtpl in the re-tagged template."""
    from docxtpl import DocxTemplate

    tpl = DocxTemplate(str(dst_path))
    variables = tpl.get_undeclared_template_variables()
    print(f"\nJinja2 variables found ({len(variables)}):")
    for var in sorted(variables):
        print(f"  {{{{ {var} }}}}")
    return variables


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Re-tag MP template placeholders to Jinja2 syntax for docxtpl"
    )
    parser.add_argument("--src", type=Path, default=DEFAULT_SRC, help="Original template path")
    parser.add_argument("--dst", type=Path, default=DEFAULT_DST, help="Re-tagged output path")
    parser.add_argument("--verify", action="store_true", help="Verify variables with docxtpl")
    args = parser.parse_args()

    if not args.src.exists():
        print(f"ERROR: Source template not found: {args.src}", file=sys.stderr)
        sys.exit(1)

    print(f"Source:  {args.src}")
    print(f"Output:  {args.dst}")
    print()

    modified_count, counters = retag(args.src, args.dst)
    print(f"Modified {modified_count} paragraphs")

    for key, count in counters.items():
        print(f"  Positional '{key}': {count} occurrence(s)")

    if args.verify:
        verify(args.dst)

    print(f"\nDone. Re-tagged template saved to: {args.dst}")


if __name__ == "__main__":
    main()
